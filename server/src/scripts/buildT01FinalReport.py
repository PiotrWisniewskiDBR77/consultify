import hashlib
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FACTS = Path(sys.argv[1]).resolve()
OUTPUT = Path(sys.argv[2]).resolve()
data = json.loads(FACTS.read_text(encoding="utf-8"))
canonical = json.dumps(data, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
digest = hashlib.sha256(canonical.encode("utf-8")).hexdigest()

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 96, 105)


def set_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = value
        cell._tc.get_or_add_tcPr().append(OxmlElement("w:shd"))
        cell._tc.get_or_add_tcPr().find(qn("w:shd")).set(qn("w:fill"), "F2F4F7")
        for run in cell.paragraphs[0].runs:
            set_font(run, 10, True, DARK)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = str(value)
            cells[idx].paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            )
            for run in cells[idx].paragraphs[0].runs:
                set_font(run, 10)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = True
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

header = section.header.paragraphs[0]
header.text = "CONSULTIFY · TRANSFORMATION CASE"
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_font(header.runs[0], 9, True, MUTED)
footer = section.footer.paragraphs[0]
footer.text = f"Case {data['transformationCaseId']} · facts {digest[:12]}"
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(footer.runs[0], 8, False, MUTED)

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(12)
title.paragraph_format.space_after = Pt(4)
set_font(title.add_run("RAPORT Z TRANSFORMACJI"), 26, True, RGBColor(0, 0, 0))
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(14)
set_font(subtitle.add_run(data["mandate"]), 14, False, MUTED)
for label, value in (
    ("Transformation Case", data["transformationCaseId"]),
    ("Lineage", data["lineageId"]),
    ("Status", data["status"]),
    ("Wspólny facts digest", digest),
):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(f"{label}: "), 10, True)
    set_font(p.add_run(value), 10)

doc.add_heading("Wynik zarządczy", level=1)
p = doc.add_paragraph(data["executiveConclusion"])
p.paragraph_format.space_after = Pt(10)

doc.add_heading("Przebieg end-to-end", level=1)
for idx, stage in enumerate(data["stages"], 1):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(stage), 11)

doc.add_page_break()
doc.add_heading("Trwałe rezultaty procesu", level=1)
labels = {
    "ideas": "Idee",
    "interviewAssignments": "Przypisania Interview",
    "acceptedInsights": "Zaakceptowane Insights",
    "drdAssessments": "Oceny DRD",
    "initiatives": "Initiatives",
    "tasks": "Zadania",
    "milestones": "Kamienie milowe",
    "decisions": "Decyzje",
    "kpis": "Karty KPI",
    "benefits": "Korzyści",
    "verifiedBenefitMeasurements": "Zweryfikowane pomiary korzyści",
}
output_rows = [(labels[key], value) for key, value in data["outputs"].items()]
add_table(["Artefakt / rekord", "Liczba"], output_rows, [6900, 2460])

doc.add_heading("Finanse i KPI", level=1)
finance = data["finance"]
kpi = data["kpi"]
add_table(
    ["Miara", "Plan", "Wynik", "Ocena"],
    [
        ("Korzyść roczna", f"{finance['forecastBenefitAnnual']:,} {finance['currency']}", f"{finance['actualBenefitAnnual']:,} {finance['currency']}", f"{finance['actualVsForecastPct']}%"),
        (kpi["name"], f"baseline {kpi['baseline']} {kpi['unit']}", f"{kpi['sustainabilityActual']} {kpi['unit']}", kpi["status"]),
    ],
    [3000, 2200, 2200, 1960],
)

doc.add_page_break()
doc.add_heading("Governance i trwałość", level=1)
governance = data["governance"]
for label, value in (
    ("Autonomia", governance["autonomy"]),
    ("Decyzja portfelowa", governance["portfolioDecision"]),
    ("Status Initiative", governance["initiativeFinalStatus"]),
    ("Efektywność", governance["effectiveness"]),
    ("Trwałość", f"{governance['sustainabilityConclusion']} · {kpi['windowDays']} dni"),
):
    p = doc.add_paragraph()
    set_font(p.add_run(f"{label}: "), 11, True)
    set_font(p.add_run(str(value)), 11)

doc.add_heading("Lineage i pakiet dowodów", level=1)
for item in data["evidence"]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(item), 10.5)

doc.add_heading("Ograniczenia dowodowe", level=1)
for item in data["evidenceLimitations"]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run(item), 10.5, False, RGBColor(155, 28, 28))

doc.add_heading("Integralność raportu", level=1)
p = doc.add_paragraph()
set_font(p.add_run("SHA-256 faktów: "), 10, True)
set_font(p.add_run(digest), 9)
p = doc.add_paragraph("Word i PowerPoint są generowane z tego samego modelu faktów. Zmiana danych wejściowych musi utworzyć nowy digest i nową wersję obu artefaktów.")
set_font(p.runs[0], 10)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(json.dumps({"output": str(OUTPUT), "factsDigest": digest}, ensure_ascii=False))
